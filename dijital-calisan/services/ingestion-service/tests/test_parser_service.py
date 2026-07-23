import io
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from app.services.parser_service import ParserError, normalize_text, parse_document


def _text_pdf() -> bytes:
    output = io.BytesIO()
    pdf = canvas.Canvas(output)
    pdf.drawString(72, 750, "OfficeIQ yillik izin proseduru ve calisan kurallari.")
    pdf.showPage()
    pdf.drawString(72, 750, "Izin talebi yonetici onayina gonderilir.")
    pdf.save()
    return output.getvalue()


def test_normalize_text_is_deterministic():
    assert normalize_text("  OfisIQ\r\n\tçok   iyi  ") == "OfisIQ\nçok iyi"


def test_parse_utf8_txt_preserves_paragraphs():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "policy.txt")
        path.parent.mkdir()
        path.write_text("Birinci prosedür paragrafı yeterince uzundur.\n\nİkinci paragraf çalışanlar içindir.", encoding="utf-8")
        parsed = parse_document("tenant/policy.txt")
        assert [block.block_index for block in parsed.blocks] == [0, 1]
        assert parsed.language == "tr"
        assert parsed.blocks[1].char_start > parsed.blocks[0].char_end


def test_parse_docx_preserves_heading_and_table():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "guide.docx")
        path.parent.mkdir()
        document = Document()
        document.add_heading("İzin Süreci", level=1)
        document.add_paragraph("Çalışan izin talebini sistem üzerinden oluşturur.")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Onaylayan"
        table.cell(0, 1).text = "Yönetici"
        document.save(path)
        parsed = parse_document("tenant/guide.docx")
        assert parsed.blocks[0].section_title == "İzin Süreci"
        assert any("Onaylayan | Yönetici" in block.text for block in parsed.blocks)


def test_parse_pdf_preserves_page_numbers():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "policy.pdf")
        path.parent.mkdir()
        path.write_bytes(_text_pdf())
        parsed = parse_document("tenant/policy.pdf")
        assert parsed.page_count == 2
        assert [block.page_number for block in parsed.blocks] == [1, 2]


def test_scanned_pdf_is_rejected_with_stable_code():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "scan.pdf")
        path.parent.mkdir()
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        with path.open("wb") as stream:
            writer.write(stream)
        with pytest.raises(ParserError) as error:
            parse_document("tenant/scan.pdf")
        assert error.value.code == "SCANNED_PDF_UNSUPPORTED"


def test_encrypted_pdf_is_rejected_with_stable_code():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "encrypted.pdf")
        path.parent.mkdir()
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        writer.encrypt("secret")
        with path.open("wb") as stream:
            writer.write(stream)
        with pytest.raises(ParserError) as error:
            parse_document("tenant/encrypted.pdf")
        assert error.value.code == "ENCRYPTED_PDF"


def test_corrupted_pdf_is_rejected_with_stable_code():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        path = Path(root, "tenant", "broken.pdf")
        path.parent.mkdir()
        path.write_bytes(b"not a valid pdf")
        with pytest.raises(ParserError) as error:
            parse_document("tenant/broken.pdf")
        assert error.value.code == "CORRUPTED_DOCUMENT"


def test_storage_path_traversal_is_rejected():
    with tempfile.TemporaryDirectory() as root, patch("app.services.parser_service.STORAGE_PATH", root):
        with pytest.raises(ParserError) as error:
            parse_document("../../etc/passwd")
        assert error.value.code == "INVALID_STORAGE_KEY"
